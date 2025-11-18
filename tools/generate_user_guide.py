import os
import re
import sys
import subprocess
import io
from docx.shared import Pt, Inches

# Danh sách packages cần thiết
REQUIRED_PACKAGES = {
    'docx': 'python-docx',
    'flask_babel': 'flask-babel'
}

def ensure_package(import_name, package_name):
    """Kiểm tra và cài đặt package nếu thiếu"""
    try:
        __import__(import_name)
    except ModuleNotFoundError:
        print(f"Module '{package_name}' chưa được cài đặt. Đang thử cài đặt bằng pip...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
            print(f"✅ Đã cài đặt {package_name}")
        except Exception as e:
            print(f"Tự động cài đặt thất bại: {e}")
            print("Hãy kích hoạt virtualenv rồi chạy thủ công:")
            print(f"  {sys.executable} -m pip install {package_name}")
            sys.exit(1)

# Kiểm tra và cài đặt tất cả packages cần thiết
for import_name, package_name in REQUIRED_PACKAGES.items():
    ensure_package(import_name, package_name)

# Import sau khi đã đảm bảo packages tồn tại
from docx import Document
from docx.shared import Pt

OUT_DIR = r"e:\C-cash_bks_repo\docs"
# Preferred source files (Vietnamese first then English sections already preserved in each md)
MD_ORDER = [
    os.path.join(OUT_DIR, "login_guide.md"),
    os.path.join(OUT_DIR, "user_panel_guide.md"),
    os.path.join(OUT_DIR, "admin_panel_guide.md"),
    os.path.join(OUT_DIR, "sftp_guide.md"),
    os.path.join(OUT_DIR, "C-cash_bks_repo.docs"),
    os.path.join(OUT_DIR, "C-cash_bks_repo.md"),
]
OUT_FILE = os.path.join(OUT_DIR, "C-cash_bks_TechSpec.docx")

def ensure_outdir():
    os.makedirs(OUT_DIR, exist_ok=True)

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    run = p.add_run("".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)

def add_list_item(doc, text, level):
    # simple bullet with indent
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(12 * level)
    p.add_run("• " + text)

def _gather_markdowns(md_list):
    """Return path to merged temporary markdown combining available md files in order."""
    found = []
    for p in md_list:
        if os.path.exists(p):
            found.append(p)
    if not found:
        return None
    merged = os.path.join(OUT_DIR, "C-cash_bks_TechSpec_merged.md")
    with open(merged, "w", encoding="utf-8") as out:
        for src in found:
            out.write(f"\n\n<!-- source: {os.path.basename(src)} -->\n\n")
            with open(src, "r", encoding="utf-8") as inf:
                out.write(inf.read())
    return merged

def _add_image_safe(doc, img_path):
    """Try to add image to doc if exists; ignore otherwise."""
    try:
        # scale down large images to max width 6 inches for doc readability
        doc.add_picture(img_path, width=Inches(6))
    except Exception as e:
        print(f"Warning: cannot add image {img_path}: {e}")

def md_to_docx(md_path, out_path):
    doc = Document()
    doc.core_properties.title = "C-cash_bks_repo - User Guide"
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_buf = []
    for raw in lines:
        line = raw.rstrip("\n")
        # image markdown: ![alt](path)
        mimg = re.match(r'!\[.*?\]\((.+?)\)', line.strip())
        if mimg:
            img_rel = mimg.group(1).strip()
            img_path = img_rel
            if not os.path.isabs(img_path):
                candidate = os.path.join(OUT_DIR, img_rel)
                if os.path.exists(candidate):
                    img_path = candidate
                else:
                    candidate2 = os.path.join(os.path.dirname(md_path), img_rel)
                    if os.path.exists(candidate2):
                        img_path = candidate2
            if os.path.exists(img_path):
                _add_image_safe(doc, img_path)
            else:
                print("Image not found:", img_rel)
            continue
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, [ln + "\n" for ln in code_buf])
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            else:
                # smaller headings -> bold paragraph
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            continue

        # Lists (unordered)
        m = re.match(r'^\s*[-\*\+]\s+(.*)', line)
        if m:
            indent = (len(line) - len(line.lstrip())) // 4
            add_list_item(doc, m.group(1).strip(), indent)
            continue

        # numbered lists
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            indent = (len(line) - len(line.lstrip())) // 4
            add_list_item(doc, m.group(1).strip(), indent)
            continue

        # blank line -> paragraph break
        if line.strip() == "":
            continue

        # normal paragraph
        p = doc.add_paragraph(line.strip())
    doc.save(out_path)

if __name__ == "__main__":
    ensure_outdir()
    # build merged markdown from ordered list
    merged_md = _gather_markdowns(MD_ORDER)
    if not merged_md:
        print("Không tìm thấy file markdown đầu vào trong docs/. Hãy kiểm tra các file .md.")
        sys.exit(1)
    md_to_docx(merged_md, OUT_FILE)
    print("Đã tạo:", OUT_FILE)
